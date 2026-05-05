import re
from dataclasses import dataclass
from pathlib import Path
from uuid import UUID

import docx
import pdfplumber
from langchain_core.documents import Document

from app.infrastructure.llm import llm_client
from app.models.schemas import ParsedResumeData

RESUME_PARSE_SYSTEM_PROMPT = """
你是一个简历结构化提取助手。
请从简历原文中提取候选人信息，并严格输出 JSON。
无法确定的字段返回 null、空字符串或空数组，不要臆造。
"""


@dataclass
class ResumeParseResult:
    # dataclass 用来打包解析结果，比返回 tuple 更清楚：
    # raw_text 是原文，parsed_data 是结构化字段，documents 是要写入向量库的文本块。
    raw_text: str
    parsed_data: ParsedResumeData
    documents: list[Document]


class ResumeParser:
    async def parse(self, file_path: Path, file_type: str, resume_id: UUID, position_id: UUID) -> ResumeParseResult:
        # 一份简历进入系统后的核心处理流程：
        # 提取文本 -> 清洗文本 -> LLM 结构化解析 -> 构造向量检索文档。
        raw_text = self.sanitize_text(await self.extract_text(file_path, file_type))
        parsed = await llm_client.complete_json(
            RESUME_PARSE_SYSTEM_PROMPT,
            f"简历原文如下：\n{raw_text}",
            ParsedResumeData,
        )
        documents = self.build_documents(parsed, resume_id, position_id)
        return ResumeParseResult(raw_text=raw_text, parsed_data=parsed, documents=documents)

    async def extract_text(self, file_path: Path, file_type: str) -> str:
        # PDF 和 DOCX 的底层格式不同，所以使用不同库提取文本。
        # 提取结果只是普通字符串，后续统一交给 LLM 解析。
        if file_type == "pdf":
            with pdfplumber.open(file_path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
        if file_type == "docx":
            document = docx.Document(file_path)
            return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
        raise ValueError("Unsupported file type")

    def sanitize_text(self, text: str) -> str:
        # 简历导出工具可能产生 NUL 字节和不可见控制字符，这些字符容易导致数据库写入失败。
        cleaned = text.replace("\x00", "")
        cleaned = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", "", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
        return cleaned.strip()

    def build_documents(
        self,
        data: ParsedResumeData,
        resume_id: UUID,
        position_id: UUID,
    ) -> list[Document]:
        # 向量库不是存整份简历 PDF，而是存“适合检索的文本块”。
        # 第一块是候选人摘要和技能，后面每段工作经历单独成块，召回会更精确。
        docs = [
            Document(
                page_content=(
                    f"{data.name} {data.job_intention or ''} {data.summary} "
                    f"技能 {' '.join(data.skills)}"
                ).strip(),
                metadata={
                    "resume_id": str(resume_id),
                    "position_id": str(position_id),
                    "doc_type": "summary",
                },
            )
        ]
        for exp in data.work_experience:
            docs.append(
                Document(
                    page_content="\n".join(
                        [f"{exp.company} {exp.position} {exp.duration}"]
                        + exp.responsibilities
                        + exp.achievements
                    ).strip(),
                    metadata={
                        "resume_id": str(resume_id),
                        "position_id": str(position_id),
                        "doc_type": "experience",
                        "company": exp.company,
                    },
                )
            )
        return docs


resume_parser = ResumeParser()
